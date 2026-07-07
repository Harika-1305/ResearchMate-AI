from docx import Document


def create_report(
    summary,
    gaps,
    questions,
    ppt
):

    doc = Document()

    doc.add_heading(
        "Research Report",
        level=1
    )

    doc.add_heading(
        "Summary",
        level=2
    )

    doc.add_paragraph(
        summary
    )

    doc.add_heading(
        "Research Gaps",
        level=2
    )

    doc.add_paragraph(
        gaps
    )

    doc.add_heading(
        "Interview Questions",
        level=2
    )

    doc.add_paragraph(
        questions
    )

    doc.add_heading(
        "PPT Outline",
        level=2
    )

    doc.add_paragraph(
        ppt
    )

    report_path = (
        "Research_Report.docx"
    )

    doc.save(
        report_path
    )

    return report_path